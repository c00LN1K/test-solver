from collections import defaultdict

from docx import Document


def read_docx(file_path):
    doc = Document(file_path)
    table_data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        if paragraph.runs[0].underline:
                            paragraph.text = '(*)' + paragraph.text
                row_data.append(cell.text.strip(' .').replace('\n', ''))
            if len(row_data) == 3 and len(row_data[2]) > 4:
                table_data.append(row_data)

    return table_data


def make_dict(data):
    res = defaultdict(dict)
    for number_question, question, answers in data:
        res[number_question]['question'] = question
        t = {}
        for i, answer in enumerate(answers.split(';')):
            if answer.startswith('(*)'):
                res[number_question]['correct'] = f'ans{i}'
                answer = answer[3:]
            t[f'ans{i}'] = answer
        res[number_question]['answers'] = t

    return res
